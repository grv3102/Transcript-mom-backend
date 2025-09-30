from io import BytesIO
from docx import Document
from typing import Dict, Any, List
from datetime import datetime

def generate_doc(data: Dict[str, Any]) -> BytesIO:
    """Generate DOCX document from structured meeting data"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Meeting Minutes', 0)
    title.alignment = 1  # Center alignment
    
    # Metadata
    processed_at = data.get('processed_at', datetime.utcnow().isoformat())
    model_used = data.get('model_used', 'AI Assistant')
    
    doc.add_paragraph(f"Generated on: {processed_at}")
    doc.add_paragraph(f"Processed by: {model_used}")
    doc.add_paragraph("")  # Empty line
    
    # Participants
    participants = data.get('participants', [])
    if participants:
        doc.add_heading('Participants', level=1)
        for participant in participants:
            doc.add_paragraph(f"• {participant}")
        doc.add_paragraph("")  # Empty line
    
    # Summary
    summary = data.get('summary', [])
    if summary:
        doc.add_heading('Summary', level=1)
        if isinstance(summary, list):
            for point in summary:
                doc.add_paragraph(f"• {point}")
        else:
            doc.add_paragraph(str(summary))
        doc.add_paragraph("")  # Empty line
    
    # Agenda Items
    agenda = data.get('agenda', [])
    if agenda:
        doc.add_heading('Agenda Items', level=1)
        for item in agenda:
            doc.add_paragraph(f"• {item}")
        doc.add_paragraph("")  # Empty line
    
    # Topics Discussed
    topics = data.get('topics', [])
    if topics:
        doc.add_heading('Topics Discussed', level=1)
        for topic in topics:
            if isinstance(topic, dict):
                topic_name = topic.get('topic', '')
                confidence = topic.get('confidence', 0)
                doc.add_paragraph(f"• {topic_name} (Confidence: {confidence:.1%})")
            else:
                doc.add_paragraph(f"• {topic}")
        doc.add_paragraph("")  # Empty line
    
    # Decisions Made
    decisions = data.get('decisions', [])
    if decisions:
        doc.add_heading('Decisions Made', level=1)
        for decision in decisions:
            doc.add_paragraph(f"• {decision}")
        doc.add_paragraph("")  # Empty line
    
    # Action Items
    action_items = data.get('actionItems', [])
    if action_items:
        doc.add_heading('Action Items', level=1)
        
        # Create table for action items
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Task'
        hdr_cells[1].text = 'Owner'
        hdr_cells[2].text = 'Deadline'
        hdr_cells[3].text = 'Status'
        
        # Data rows
        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('task', ''))
            row_cells[1].text = str(item.get('owner', ''))
            row_cells[2].text = str(item.get('deadline', 'Not specified'))
            row_cells[3].text = str(item.get('status', 'Pending'))
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
